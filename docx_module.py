from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

template_path = 'templates/fiz_p.docx'

replacements = {
    '%CURDATE%': {'text': '', 'format': 3},
    '%ISTNAME%': {'text': '', 'format': 1},
    '%ISTINN%': {'text': '', 'format': 1},
    '%ISTADDRESS%': {'text': '', 'format': 2},
    '%OTVNAME%': {'text': '', 'format': 1},
    '%OTVADDRESS%': {'text': '', 'format': 2},
    '%OTVBD%': {'text': '', 'format': 2},
    '%OTVCITY%': {'text': '', 'format': 1},
    '%OTVPS%': {'text': '', 'format': 2},
    '%OTVPN%': {'text': '', 'format': 2},
    '%OTVPASSDATE%': {'text': '', 'format': 2},
    '%UCHDATE%': {'text': '', 'format': 1},
    '%UCHNUM%': {'text': '', 'format': 1},
    '%UCHKADASTR%': {'text': '', 'format': 1},
    '%USCHSQ%': {'text': '', 'format': 1},
    '%UCHADDRESS%': {'text': '', 'format': 1},
    '%UCHPREDN%': {'text': '', 'format': 1},
    '%LONG%': {'text': '', 'format': 1},
    '%ARSUM%': {'text': '', 'format': 1},
    '%DOLGSTART%': {'text': '', 'format': 1},
    '%DOLGEND%': {'text': '', 'format': 1},
    '%DSUM%': {'text': '', 'format': 1},
    '%PSUM%': {'text': '', 'format': 1},
    '%DOLGSUM%': {'text': '', 'format': 1},
    '%PRDATE%': {'text': '', 'format': 1},
    '%PRNUM%': {'text': '', 'format': 1},
    '%OVTBD%': {'text': '', 'format': 1}
}


'''
Расшифровка меток в шаблоне fiz_p (Ответчик - физ. лицо + пени)
1. %CURDATE% - Текущая дата
2. %ISTNAME% - Название истца
3. %ISTINN% - ИНН истца
4. %ISTADDRESS% - Адрес истца
5. %OTVNAME% - Имя ответчика
6. %OTVADDRESS% - Адрес ответчика
7. %OTVBD% - Дата рождения ответчика
8. %OTVCITY% - Город рождения ответчика
9. %OTVPS% - Серия паспорта ответчика
10. %OTVPN% - Номер паспорта ответчика
11. %OTVPASSDATE% - Дата выдачи паспорта ответчика
12. %UCHDATE% - Дата заключения договора
13. %UCHNUM% - Номер договора
14. %UCHKADASTR% - Кадастровый номер земельного участка
15. %USCHSQ% - Площадь земельного участка
16. %UCHADDRESS% - Адрес земельного участка
17. %UCHPREDN% - Предназначение земельного участка
18. %LONG% - Срок действия договора
19. %ARSUM% - Размер арендной платы
20. %DOLGSTART% - Дата начала задолженности
21. %DOLGEND% - Дата окончания задолженности
22. %DSUM% - Сумма задолженности
23. %PSUM% - Сумма пени
24. %DOLGSUM% - Общая сумма задолженности
25. %PRDATE% - Дата претензии
26. %PRNUM% - Номер претензии
27. %OVTBD% - Дата рождения ответчика
'''


# Формат: 1 - обычный 12, 2 - обычный 10, 3 - полужирный 10,
def replace_tag(struct, tag, text, format):
    if format == 1:
        text = text.replace('\n', ' ')
        struct.text = struct.text.replace(tag, text)
        for run in struct.runs:
            run.font.name = 'PT Astra Serif'
            run.font.size = Pt(12)
    elif format == 2:
        text = text.replace('\n', ' ')
        struct.text = struct.text.replace(tag, text)
        for paragraph in struct.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.name = 'PT Astra Serif'
                run.font.size = Pt(10)
    elif format == 3:
        text = text.replace('\n', ' ')
        struct.text = struct.text.replace(tag, text)
        for paragraph in struct.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'PT Astra Serif'
                run.font.size = Pt(10)


# Поиск в шапке
def find_labels_in_header(table):
    current_date = datetime.now().strftime('%d.%m.%Y')

    for row in table.rows:
        for cell in row.cells:
            if '%ISTADDRESS%' in cell.text:
                replace_tag(cell, '%ISTADDRESS%', 'ул. Пушкина', replacements['%ISTADDRESS%']['format'])
            if '%CURDATE%' in cell.text:
                replace_tag(cell, '%CURDATE%', current_date, 3)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = 'PT Astra Serif'
                        run.font.size = Pt(10)


def find_labels_in_document(doc):
    for table in doc.tables:
        find_labels_in_header(table)


# Загружаем документ
doc_path = 'templates/fiz_p.docx'
doc = Document(doc_path)

# Вызываем функцию для поиска меток в документе
find_labels_in_document(doc)


# Замена текста в параграфах
def replace_paragraphs(doc, placeholder, text):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            replace_tag(paragraph, placeholder, text, replacements[placeholder]['format'])
            for run in paragraph.runs:
                run.font.name = 'PT Astra Serif'
                run.font.size = Pt(12)
            # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


replace_paragraphs(doc, '%ISTADDRESS%', 'aa')

doc.save('output.docx')
